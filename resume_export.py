"""
简历 Word 导出模块
排版风格复刻 RenderCV Classic 主题：
- 蓝色 Section 标题 + 右侧蓝色横线
- 每个条目使用双栏表格：左侧标题+要点，右侧日期（右对齐）
- 紧凑行距，专业排版
"""

import io
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── 样式常量（RenderCV Classic）─────────────────────────
BRAND_BLUE = RGBColor(0, 79, 144)
TEXT_BLACK = RGBColor(0, 0, 0)
TEXT_DARK = RGBColor(51, 51, 51)
GRAY = RGBColor(110, 110, 110)
LIGHT_GRAY = RGBColor(180, 180, 180)

FONT_BODY = 'Calibri'
FONT_CN = '微软雅黑'

SIZE_BODY = Pt(10)
SIZE_NAME = Pt(24)
SIZE_SECTION = Pt(12)
SIZE_ENTRY_TITLE = Pt(10)
SIZE_DATE = Pt(9.5)
SIZE_BULLET = Pt(10)

PAGE_MARGIN = Inches(0.7)
LINE_SPACING = Pt(14)
BULLET_LINE_SPACING = Pt(13.5)
DATE_COL_WIDTH = Cm(4.5)

SECTION_META = {
    'education':   (0, '教育经历'),
    'internship':  (1, '实习经历'),
    'project':     (2, '项目经历'),
    'competition': (3, '竞赛经历'),
    'skill':       (4, '技能证书'),
    'other':       (5, '其他'),
}

# 日期正则 — 匹配完整的日期范围（包含"至今"等结尾词）
_DATE_RANGE_RE = re.compile(
    r'\d{4}\s*[年./\-]\s*\d{1,2}\s*[月]?\s*[-–—~至到]+\s*'
    r'(?:\d{4}\s*[年./\-]\s*\d{1,2}\s*[月]?|至今|今|现在|present)',
    re.IGNORECASE
)
# 单独日期
_DATE_SINGLE_RE = re.compile(
    r'\d{4}\s*[年./\-]\s*\d{1,2}\s*[月]?',
    re.IGNORECASE
)


# ═══════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════

def _set_font(run, size=SIZE_BODY, bold=False, color=TEXT_BLACK, italic=False):
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.name = FONT_BODY
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)


def _set_spacing(para, before=Pt(0), after=Pt(0), line=LINE_SPACING):
    pf = para.paragraph_format
    pf.space_before = before
    pf.space_after = after
    if line:
        pf.line_spacing = line


def _no_borders(tbl_element):
    tblPr = tbl_element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl_element.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>'))
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        '<w:top w:w="0" w:type="dxa"/>'
        '<w:left w:w="0" w:type="dxa"/>'
        '<w:bottom w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/>'
        '</w:tblCellMar>'
    ))


def _set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:type="dxa" w:w="{int(width.emu / 635)}"/>'))


def _cell_valign_top(cell):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))


def _add_bottom_border(para, color="004F90", sz="6"):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._p.insert(0, pPr)
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))


# ═══════════════════════════════════════════
# 内容预处理 + 解析
# ═══════════════════════════════════════════

def _clean_content(raw: str) -> str:
    """
    清理简历文本（移植自前端 cleanResumeContent）：
    1. 孤立 bullet 符号合并到下一行
    2. 句中断行合并
    """
    lines = raw.split('\n')
    merged = []
    i = 0
    while i < len(lines):
        trimmed = lines[i].strip()
        # 单独的 bullet 符号 → 合并到下一行
        if re.match(r'^[•·\-*●○]$', trimmed) and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if next_line:
                merged.append(f'{trimmed} {next_line}')
                i += 2
                continue
        merged.append(lines[i])
        i += 1
    return '\n'.join(merged)


def _is_bullet(line):
    """判断是否为 bullet 行（允许 bullet 后无空格直接跟中文）"""
    return bool(re.match(r'^[•·\-*●○▪►]\s*\S', line.strip()))


def _clean_bullet_text(line):
    return re.sub(r'^[•·\-*●○▪►]\s*', '', line.strip())


def _extract_date_from_text(text: str) -> str:
    """从文本中提取日期范围，优先匹配完整范围"""
    m = _DATE_RANGE_RE.search(text)
    if m:
        return m.group(0).strip()
    m = _DATE_SINGLE_RE.search(text)
    if m:
        return m.group(0).strip()
    return ''


def _is_mostly_date(line: str) -> bool:
    """判断一行是否主要由日期组成"""
    stripped = line.strip()
    if not stripped:
        return False
    date_str = _extract_date_from_text(stripped)
    if not date_str:
        return False
    # 去掉日期后剩余的有效字符
    remaining = stripped
    for pattern in [_DATE_RANGE_RE, _DATE_SINGLE_RE]:
        remaining = pattern.sub('', remaining)
    remaining = re.sub(r'[\s\-–—~至到月年/.|,，]+', '', remaining)
    return len(remaining) < max(len(stripped) * 0.35, 5)


def _deduplicate_content(title: str, lines: list) -> list:
    """去除 content 中与 title 重复的行"""
    if not title:
        return lines
    # 标题拆成关键词（"弗若斯特沙利文咨询-行业分析实习生" → ["弗若斯特沙利文咨询", "行业分析实习生"]）
    title_parts = re.split(r'[-–—/|·,，]', title)
    title_parts = [p.strip() for p in title_parts if len(p.strip()) >= 2]

    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append(line)
            continue
        # 如果这一行完全是标题的某个部分 → 跳过
        is_dup = False
        for part in title_parts:
            if stripped == part:
                is_dup = True
                break
        if not is_dup:
            result.append(line)
    return result


def _remove_title_prefix(text: str, title: str) -> str:
    """从文本开头去除与 title 重复的部分"""
    if not title or not text:
        return text
    title_parts = re.split(r'[-–—/|·,，]', title)
    title_parts = [p.strip() for p in title_parts if len(p.strip()) >= 2]
    result = text
    for part in title_parts:
        if result.startswith(part):
            result = result[len(part):].strip()
            result = re.sub(r'^[-–—/|·,，\s]+', '', result).strip()
    return result


def _parse_entry(title: str, content: str):
    """
    解析条目内容，返回：
    - date_text: 提取出的日期
    - sub_entries: [{subtitle, bullets}] 子条目列表
    """
    content = _clean_content(content)
    lines = content.split('\n')
    lines = _deduplicate_content(title, lines)

    date_text = ''
    sub_entries = []
    current_subtitle = ''
    current_bullets = []

    def flush():
        nonlocal current_subtitle, current_bullets
        if current_subtitle or current_bullets:
            sub_entries.append({
                'subtitle': current_subtitle,
                'bullets': current_bullets[:],
            })
        current_subtitle = ''
        current_bullets = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 提取日期（从任何含日期的行中）
        if not date_text:
            d = _extract_date_from_text(stripped)
            if d:
                date_text = d

        # 纯日期行 → 跳过（日期已提取到右列）
        if _is_mostly_date(stripped):
            continue

        # bullet 行
        if _is_bullet(stripped):
            current_bullets.append(_clean_bullet_text(stripped))
            continue

        # 非 bullet 行：可能是子标题（如项目名）
        # 先清理日期残余和标题重复部分
        clean_line = _DATE_RANGE_RE.sub('', stripped)
        clean_line = _DATE_SINGLE_RE.sub('', clean_line)
        clean_line = re.sub(r'^\s*[-–—~至到,，|]+\s*', '', clean_line).strip()
        clean_line = re.sub(r'\s*[-–—~至到,，|]+\s*$', '', clean_line).strip()
        clean_line = _remove_title_prefix(clean_line, title)
        if not clean_line or len(clean_line) < 2:
            continue

        # 如果当前已经有 bullets → 说明上一段结束了，开始新子条目
        if current_bullets:
            flush()
            current_subtitle = clean_line
        elif not current_subtitle:
            current_subtitle = clean_line
        else:
            current_subtitle += '  ' + clean_line

    flush()
    return date_text, sub_entries


# ═══════════════════════════════════════════
# Section 标题
# ═══════════════════════════════════════════

def _add_section_heading(doc, label):
    table = doc.add_table(rows=1, cols=2)
    _no_borders(table._tbl)

    # 左列：蓝色粗体标题
    cell_l = table.cell(0, 0)
    p = cell_l.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p, before=Pt(16), after=Pt(3), line=None)
    run = p.add_run(label)
    _set_font(run, size=SIZE_SECTION, bold=True, color=BRAND_BLUE)
    # 不换行 + 自适应宽度
    tc_l = cell_l._tc
    tcPr_l = tc_l.find(qn('w:tcPr'))
    if tcPr_l is None:
        tcPr_l = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc_l.insert(0, tcPr_l)
    tcPr_l.append(parse_xml(f'<w:noWrap {nsdecls("w")}/>'))
    tcPr_l.append(parse_xml(f'<w:tcW {nsdecls("w")} w:type="auto" w:w="0"/>'))

    # 右列：蓝色横线
    cell_r = table.cell(0, 1)
    p2 = cell_r.paragraphs[0]
    _set_spacing(p2, before=Pt(16), after=Pt(3), line=None)
    _add_bottom_border(p2)
    run2 = p2.add_run(' ')
    _set_font(run2, size=SIZE_SECTION, color=BRAND_BLUE)


# ═══════════════════════════════════════════
# 条目渲染
# ═══════════════════════════════════════════

def _add_entry(doc, title: str, content: str):
    """
    RenderCV Classic 风格双栏条目：
    ┌─────────────────────────────────┬────────────┐
    │ **Title**                       │ Date Range │
    │ Subtitle (if any)               │            │
    │ • bullet 1                      │            │
    │ • bullet 2                      │            │
    │ Sub-project title               │            │
    │ • bullet 3                      │            │
    └─────────────────────────────────┴────────────┘
    """
    date_text, sub_entries = _parse_entry(title, content)

    table = doc.add_table(rows=1, cols=2)
    _no_borders(table._tbl)

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    _set_cell_width(cell_right, DATE_COL_WIDTH)
    _cell_valign_top(cell_left)
    _cell_valign_top(cell_right)

    # ── 左列 ──────────────────────────
    # 主标题（粗体）
    p_title = cell_left.paragraphs[0]
    _set_spacing(p_title, before=Pt(7), after=Pt(1), line=LINE_SPACING)
    if title:
        run_t = p_title.add_run(title)
        _set_font(run_t, size=SIZE_ENTRY_TITLE, bold=True, color=TEXT_BLACK)

    # 子条目
    for si, sub in enumerate(sub_entries):
        # 子标题（如果有，且跟主标题不同）
        if sub['subtitle'] and sub['subtitle'] != title:
            p_sub = cell_left.add_paragraph()
            _set_spacing(p_sub, before=Pt(4) if si > 0 else Pt(1), after=Pt(1), line=LINE_SPACING)
            run_sub = p_sub.add_run(sub['subtitle'])
            _set_font(run_sub, size=SIZE_BODY, color=TEXT_DARK, italic=True)

        # Bullet 要点
        for b in sub['bullets']:
            p_b = cell_left.add_paragraph()
            _set_spacing(p_b, before=Pt(0), after=Pt(0.5), line=BULLET_LINE_SPACING)
            p_b.paragraph_format.left_indent = Cm(0.35)
            p_b.paragraph_format.first_line_indent = Cm(-0.35)
            run_dot = p_b.add_run('•  ')
            _set_font(run_dot, size=SIZE_BULLET, color=GRAY)
            run_text = p_b.add_run(b)
            _set_font(run_text, size=SIZE_BULLET, color=TEXT_DARK)

    # ── 右列：日期 ────────────────────
    p_date = cell_right.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_spacing(p_date, before=Pt(7), after=Pt(0), line=LINE_SPACING)
    if date_text:
        run_d = p_date.add_run(date_text)
        _set_font(run_d, size=SIZE_DATE, color=GRAY)


def _add_skill_entry(doc, title: str, content: str):
    """技能条目：Label: Details"""
    content = _clean_content(content)
    lines = [l.strip() for l in content.split('\n') if l.strip()]

    para = doc.add_paragraph()
    _set_spacing(para, before=Pt(5), after=Pt(2), line=LINE_SPACING)
    if title:
        run_label = para.add_run(title + '：')
        _set_font(run_label, size=SIZE_BODY, bold=True, color=TEXT_BLACK)
    if lines:
        run_detail = para.add_run('  '.join(lines))
        _set_font(run_detail, size=SIZE_BODY, color=TEXT_DARK)


# ═══════════════════════════════════════════
# 主函数
# ═══════════════════════════════════════════

def generate_resume_docx(resume_sections: list, user_name: str = '') -> io.BytesIO:
    doc = Document()

    # 页面设置
    sec = doc.sections[0]
    sec.top_margin = PAGE_MARGIN
    sec.bottom_margin = PAGE_MARGIN
    sec.left_margin = PAGE_MARGIN
    sec.right_margin = PAGE_MARGIN

    # 姓名
    if user_name:
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p_name, before=Pt(0), after=Pt(12), line=None)
        run = p_name.add_run(user_name)
        _set_font(run, size=SIZE_NAME, bold=True, color=BRAND_BLUE)

    # 排序
    sorted_sections = sorted(
        resume_sections,
        key=lambda s: SECTION_META.get(s.get('type', 'other'), (99, ''))[0]
    )

    # 渲染
    current_type = None
    for entry in sorted_sections:
        sec_type = entry.get('type', 'other')
        title = entry.get('title', '')
        content = entry.get('content', '')

        if sec_type != current_type:
            current_type = sec_type
            _, label = SECTION_META.get(sec_type, (99, sec_type))
            _add_section_heading(doc, label)

        if sec_type == 'skill':
            _add_skill_entry(doc, title, content)
        else:
            _add_entry(doc, title, content)

    # 页脚
    footer = sec.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('由校园人才估值平台生成')
    _set_font(run, size=Pt(8), color=LIGHT_GRAY)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_resume_pdf(resume_sections: list, user_name: str = '') -> io.BytesIO:
    """
    生成 PDF 版简历

    策略：先生成 Word → 用 libreoffice 转换为 PDF
    如果 libreoffice 不可用，回退到简单的文本 PDF（使用 reportlab 或纯文本）
    """
    import subprocess
    import tempfile
    import os

    # 第一步：生成 Word
    docx_buf = generate_resume_docx(resume_sections, user_name)

    # 第二步：尝试用 libreoffice 转换
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'resume.docx')
            pdf_path = os.path.join(tmpdir, 'resume.pdf')

            with open(docx_path, 'wb') as f:
                f.write(docx_buf.read())
            docx_buf.seek(0)

            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', tmpdir, docx_path],
                capture_output=True, timeout=30,
            )

            if result.returncode == 0 and os.path.exists(pdf_path):
                pdf_buf = io.BytesIO()
                with open(pdf_path, 'rb') as f:
                    pdf_buf.write(f.read())
                pdf_buf.seek(0)
                return pdf_buf

    except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
        print(f"[ResumeExport] libreoffice 转换失败: {e}，回退到纯文本 PDF")

    # 回退方案：生成简单的纯文本 PDF
    return _generate_simple_pdf(resume_sections, user_name)


def _generate_simple_pdf(resume_sections: list, user_name: str = '') -> io.BytesIO:
    """
    回退方案：用纯 Python 生成简单 PDF（不依赖 libreoffice）
    使用 fpdf2（轻量级，纯 Python）
    """
    try:
        from fpdf import FPDF
    except ImportError:
        # 如果 fpdf2 也没有，返回 Word 文件并标注
        print("[ResumeExport] fpdf2 未安装，无法生成 PDF，返回 Word")
        return generate_resume_docx(resume_sections, user_name)

    pdf = FPDF()
    pdf.add_page()

    # 尝试添加中文字体
    font_added = False
    for font_path in [
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
    ]:
        if os.path.exists(font_path):
            try:
                pdf.add_font('CJK', '', font_path, uni=True)
                pdf.set_font('CJK', size=10)
                font_added = True
                break
            except Exception:
                continue

    if not font_added:
        pdf.set_font('Helvetica', size=10)

    # 标题
    if user_name:
        pdf.set_font_size(20)
        pdf.cell(0, 15, user_name, ln=True, align='C')
        pdf.ln(5)

    # 按 section 排列
    sorted_sections = sorted(
        resume_sections,
        key=lambda s: SECTION_META.get(s.get('type', 'other'), (99, ''))[0]
    )

    current_type = None
    for entry in sorted_sections:
        sec_type = entry.get('type', 'other')
        title = entry.get('title', '')
        content = entry.get('content', '')

        if sec_type != current_type:
            current_type = sec_type
            _, label = SECTION_META.get(sec_type, (99, sec_type))
            pdf.ln(3)
            pdf.set_font_size(13)
            pdf.cell(0, 8, label, ln=True)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 170, pdf.get_y())
            pdf.ln(2)

        pdf.set_font_size(10)
        if title:
            pdf.set_font(style='B')
            pdf.cell(0, 6, title, ln=True)
            pdf.set_font(style='')

        if content:
            for line in content.split('\n'):
                stripped = line.strip()
                if stripped:
                    pdf.multi_cell(0, 5, '  ' + stripped)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf
