"""
===========================================
学生版校招身价评估 API 服务
===========================================
为学生版"校招身价测评"提供后端接口

启动方式: python mini_api.py
默认端口: 5001
"""

from flask import Flask, request, jsonify, Response, send_from_directory, send_file
from flask_cors import CORS
from config import config
from llm_service import LLMService
from chat_agent import ChatAgent
from model_router import ModelRouter, UsageTracker
from incremental_convergence import IncrementalConvergence
from validation_rules import validation_rules
from salary_calculator import SalaryCalculator
from ability_mapper import map_factors_to_dimensions, get_dimension_radar_data, get_dimension_summary
from calculator import calculate_hay_evaluation
from level_tags import get_level_tag_and_desc
from salary_competitiveness import calculate_salary_competitiveness
from school_mapper import identify_school_tier
from student_coefficients import apply_student_coefficients, format_salary_k

import traceback
import base64
import re
import io
import os
import json
import threading
from datetime import datetime

# Supabase 客户端（用户数据持久化）
supabase_client = None
try:
    _sb_url = os.getenv('SUPABASE_URL', '')
    _sb_key = os.getenv('SUPABASE_SERVICE_KEY', '')
    if _sb_url and _sb_key:
        from supabase import create_client
        supabase_client = create_client(_sb_url, _sb_key)
        print("✓ Supabase 客户端初始化成功")
    else:
        print("⚠ SUPABASE_URL/SUPABASE_SERVICE_KEY 未设置，Supabase 功能不可用")
except Exception as _sb_err:
    print(f"⚠ Supabase 初始化失败: {_sb_err}")

# 全局评估计数器 & 日志存储（PostgreSQL 持久化）
import psycopg2
import psycopg2.extras

DATABASE_URL = os.getenv('DATABASE_URL', '')

def _init_db():
    """初始化数据库表"""
    if not DATABASE_URL:
        print("[日志] ⚠️ 未设置 DATABASE_URL，日志将不会持久化")
        return
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor()
        cur.execute('''
            CREATE TABLE IF NOT EXISTS assessment_logs (
                id SERIAL PRIMARY KEY,
                time TIMESTAMP NOT NULL DEFAULT NOW(),
                type VARCHAR(20) NOT NULL,
                city VARCHAR(100),
                industry VARCHAR(100),
                function VARCHAR(100),
                title VARCHAR(200),
                resume_len INTEGER,
                retry INTEGER,
                insufficient BOOLEAN,
                resume_text TEXT,
                grade INTEGER,
                tag VARCHAR(100),
                salary VARCHAR(100),
                factors_detail JSONB,
                abilities TEXT,
                deep_insight TEXT,
                error TEXT,
                elapsed VARCHAR(20)
            )
        ''')
        # 在 assessment_logs 上追加页面停留时长列 & 学生版字段（幂等）
        for col in ['invite_code VARCHAR(50)', 'welcome_s INTEGER', 'form_s INTEGER',
                     'result_s INTEGER', 'mine_s INTEGER',
                     'school_name VARCHAR(200)', 'education_level VARCHAR(50)', 'school_tier VARCHAR(50)',
                     'major VARCHAR(200)', 'company_type VARCHAR(100)', 'target_company VARCHAR(200)']:
            col_name = col.split()[0]
            cur.execute(f'''
                DO $$ BEGIN
                    ALTER TABLE assessment_logs ADD COLUMN {col};
                EXCEPTION WHEN duplicate_column THEN NULL;
                END $$;
            ''')
        conn.commit()
        cur.close()
        conn.close()
        print("[日志] ✓ 数据库表已就绪")
    except Exception as e:
        print(f"[日志] ❌ 数据库初始化失败: {e}")

def _get_counter():
    """从数据库获取当前最大 ID 作为计数器"""
    if not DATABASE_URL:
        return 0
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor()
        cur.execute('SELECT COALESCE(MAX(id), 0) FROM assessment_logs')
        counter = cur.fetchone()[0]
        cur.close()
        conn.close()
        return counter
    except Exception:
        return 0

def _calc_resume_health(abilities: dict) -> int:
    """根据 5 维能力分数计算简历健康度（0-100）"""
    if not abilities:
        return 30
    scores = [v.get('score', 0) for v in abilities.values() if isinstance(v, dict)]
    if not scores:
        return 30
    avg = sum(scores) / len(scores)
    # 将平均分映射到健康度：avg 0-100 → health 10-95
    return max(10, min(95, int(avg * 0.85 + 10)))


def _insert_log(log_data):
    """插入一条日志到数据库，返回新 ID"""
    if not DATABASE_URL:
        return log_data.get('id', 0)
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO assessment_logs (time, type, city, industry, function, title,
                resume_len, retry, insufficient, resume_text, grade, tag, salary,
                factors_detail, abilities, deep_insight, error, elapsed,
                invite_code, welcome_s, form_s,
                school_name, education_level, school_tier, major, company_type, target_company)
            VALUES (%(time)s, %(type)s, %(city)s, %(industry)s, %(function)s, %(title)s,
                %(resume_len)s, %(retry)s, %(insufficient)s, %(resume_text)s, %(grade)s,
                %(tag)s, %(salary)s, %(factors_detail)s, %(abilities)s, %(deep_insight)s,
                %(error)s, %(elapsed)s,
                %(invite_code)s, %(welcome_s)s, %(form_s)s,
                %(school_name)s, %(education_level)s, %(school_tier)s, %(major)s, %(company_type)s, %(target_company)s)
            RETURNING id
        ''', {
            'time': log_data.get('time'),
            'type': log_data.get('type'),
            'city': log_data.get('city'),
            'industry': log_data.get('industry'),
            'function': log_data.get('function'),
            'title': log_data.get('title'),
            'resume_len': log_data.get('resume_len'),
            'retry': log_data.get('retry'),
            'insufficient': log_data.get('insufficient'),
            'resume_text': log_data.get('resume_text'),
            'grade': log_data.get('grade'),
            'tag': log_data.get('tag'),
            'salary': log_data.get('salary'),
            'factors_detail': json.dumps(log_data['factors_detail'], ensure_ascii=False) if log_data.get('factors_detail') else None,
            'abilities': log_data.get('abilities'),
            'deep_insight': log_data.get('deep_insight'),
            'error': log_data.get('error'),
            'elapsed': log_data.get('elapsed'),
            'invite_code': log_data.get('invite_code'),
            'welcome_s': log_data.get('welcome_s'),
            'form_s': log_data.get('form_s'),
            'school_name': log_data.get('school_name'),
            'education_level': log_data.get('education_level'),
            'school_tier': log_data.get('school_tier'),
            'major': log_data.get('major'),
            'company_type': log_data.get('company_type'),
            'target_company': log_data.get('target_company'),
        })
        new_id = cur.fetchone()[0]
        conn.commit()
        cur.close()
        conn.close()
        return new_id
    except Exception as e:
        print(f"[日志] ❌ 插入日志失败: {e}")
        return log_data.get('id', 0)

def _query_logs(after=0):
    """查询 id > after 的日志"""
    if not DATABASE_URL:
        return []
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute('SELECT * FROM assessment_logs WHERE id > %s ORDER BY id', (after,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        result = []
        for row in rows:
            log = dict(row)
            log['time'] = log['time'].strftime('%Y-%m-%d %H:%M:%S') if log.get('time') else ''
            if log.get('factors_detail') and isinstance(log['factors_detail'], str):
                log['factors_detail'] = json.loads(log['factors_detail'])
            result.append(log)
        return result
    except Exception as e:
        print(f"[日志] ❌ 查询日志失败: {e}")
        return []

_init_db()
assessment_counter = _get_counter()
print(f"[日志] 已加载计数器={assessment_counter}")

# PDF 解析
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
    print("✓ PDF解析支持已启用 (PyMuPDF)")
except ImportError:
    PDF_SUPPORT = False
    print("⚠ PDF解析不可用 (请安装 PyMuPDF: pip install PyMuPDF)")

# Word 文档解析
try:
    from docx import Document as DocxDocument
    from io import BytesIO
    DOCX_SUPPORT = True
    print("✓ Word文档解析支持已启用 (python-docx)")
except ImportError:
    DOCX_SUPPORT = False
    print("⚠ Word文档解析不可用 (请安装 python-docx: pip install python-docx)")


def parse_pdf_from_base64(base64_content: str) -> str:
    """从 Base64 编码的 PDF 中提取文本"""
    if not PDF_SUPPORT:
        return ""

    try:
        # 清理 Base64 内容（移除空白字符）
        cleaned_base64 = re.sub(r'\s+', '', base64_content)
        print(f"[PDF解析] Base64长度: {len(cleaned_base64)} 字符")
        print(f"[PDF解析] Base64前50字符: {cleaned_base64[:50]}")

        # 解码 Base64
        pdf_bytes = base64.b64decode(cleaned_base64)
        print(f"[PDF解析] 解码后字节数: {len(pdf_bytes)}")

        # 使用 PyMuPDF 解析
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text_parts = []

        for page in doc:
            text_parts.append(page.get_text())

        doc.close()

        full_text = "\n".join(text_parts).strip()
        print(f"[PDF解析] 成功提取 {len(full_text)} 字符")
        return full_text

    except Exception as e:
        import traceback
        print(f"[PDF解析] 失败: {e}")
        traceback.print_exc()
        return ""


def parse_docx_from_base64(base64_content: str) -> str:
    """从 Base64 编码的 Word 文档中提取文本"""
    if not DOCX_SUPPORT:
        return ""

    try:
        cleaned_base64 = re.sub(r'\s+', '', base64_content)
        print(f"[Word解析] Base64长度: {len(cleaned_base64)} 字符")

        docx_bytes = base64.b64decode(cleaned_base64)
        print(f"[Word解析] 解码后字节数: {len(docx_bytes)}")

        doc = DocxDocument(BytesIO(docx_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 也提取表格中的文本（简历常用表格布局）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        full_text = "\n".join(text_parts).strip()
        print(f"[Word解析] 成功提取 {len(full_text)} 字符")
        return full_text

    except Exception as e:
        import traceback
        print(f"[Word解析] 失败: {e}")
        traceback.print_exc()
        return ""


def extract_resume_text(resume_text: str) -> str:
    """
    处理简历文本，支持：
    1. 纯文本直接返回
    2. [文件：xxx.pdf] + Base64 格式的 PDF 文件
    3. [文件：xxx.docx] + Base64 格式的 Word 文件
    """
    if not resume_text:
        return ""

    # 检测是否是文件格式: [文件：xxx.pdf]\nBase64内容
    file_pattern = r'^\[文件[：:]\s*(.+?)\]\s*\n?(.+)$'
    match = re.match(file_pattern, resume_text, re.DOTALL)

    if match:
        filename = match.group(1)
        content = match.group(2).strip()
        filename_lower = filename.lower()

        print(f"[简历处理] 检测到文件上传: {filename}")

        # 判断文件类型
        if filename_lower.endswith('.pdf'):
            # PDF 文件 - 解析
            extracted = parse_pdf_from_base64(content)
            if extracted:
                return extracted
            else:
                print("[简历处理] PDF解析失败，无法提取文本")
                return ""
        elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
            # Word 文件 - 解析
            extracted = parse_docx_from_base64(content)
            if extracted:
                return extracted
            else:
                print("[简历处理] Word文档解析失败，无法提取文本")
                return ""
        else:
            # 其他文件类型（如 txt）- 尝试 Base64 解码为文本
            try:
                decoded = base64.b64decode(content).decode('utf-8')
                print(f"[简历处理] 文本文件解码成功: {len(decoded)} 字符")
                return decoded
            except:
                print("[简历处理] 无法解析文件内容")
                return content

    # 纯文本直接返回
    return resume_text

# ===========================================
# 邀请码系统
# ===========================================

# 从环境变量加载一次性邀请码（逗号分隔）
_raw_invite_codes = os.getenv('INVITE_CODES', '')
INVITE_CODES_ALL = {c.strip().upper() for c in _raw_invite_codes.split(',') if c.strip()}

# 从环境变量加载永久邀请码（无限次使用，逗号分隔）
_raw_permanent_codes = os.getenv('PERMANENT_INVITE_CODES', '')
PERMANENT_INVITE_CODES = {c.strip().upper() for c in _raw_permanent_codes.split(',') if c.strip()}

# 已使用的邀请码持久化文件
USED_CODES_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'used_invite_codes.json')
_codes_lock = threading.Lock()

def _load_used_codes() -> set:
    """从文件加载已使用的邀请码"""
    try:
        if os.path.exists(USED_CODES_FILE):
            with open(USED_CODES_FILE, 'r') as f:
                return set(json.load(f))
    except Exception as e:
        print(f"[邀请码] 加载已使用码失败: {e}")
    return set()

def _save_used_codes(used: set):
    """将已使用的邀请码持久化到文件"""
    try:
        with open(USED_CODES_FILE, 'w') as f:
            json.dump(list(used), f)
    except Exception as e:
        print(f"[邀请码] 保存已使用码失败: {e}")

USED_INVITE_CODES = _load_used_codes()

# 已消耗的邀请码（评估成功后标记，彻底失效）
CONSUMED_CODES_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'consumed_invite_codes.json')

def _load_consumed_codes() -> set:
    """从文件加载已消耗的邀请码"""
    try:
        if os.path.exists(CONSUMED_CODES_FILE):
            with open(CONSUMED_CODES_FILE, 'r') as f:
                return set(json.load(f))
    except Exception as e:
        print(f"[邀请码] 加载已消耗码失败: {e}")
    return set()

def _save_consumed_codes(consumed: set):
    """将已消耗的邀请码持久化到文件"""
    try:
        with open(CONSUMED_CODES_FILE, 'w') as f:
            json.dump(list(consumed), f)
    except Exception as e:
        print(f"[邀请码] 保存已消耗码失败: {e}")

CONSUMED_INVITE_CODES = _load_consumed_codes()

# 所有合法邀请码 = 一次性 + 永久
ALL_VALID_CODES = INVITE_CODES_ALL | PERMANENT_INVITE_CODES

print(f"[邀请码] 一次性邀请码: {INVITE_CODES_ALL}")
print(f"[邀请码] 永久邀请码: {PERMANENT_INVITE_CODES}")
print(f"[邀请码] 已使用邀请码: {USED_INVITE_CODES}")
print(f"[邀请码] 已消耗邀请码: {CONSUMED_INVITE_CODES}")
print(f"[邀请码] 一次性剩余可用: {INVITE_CODES_ALL - USED_INVITE_CODES - CONSUMED_INVITE_CODES}")


# ===========================================
# 初始化 Flask
# ===========================================

app = Flask(__name__, static_folder='static', static_url_path='/static')
app.config['JSON_AS_ASCII'] = False
CORS(app, origins='*')

# ===========================================
# 初始化服务
# ===========================================

print("\n" + "=" * 60)
print("学生版校招身价评估后端服务启动中...")
print("=" * 60)

# 模型路由器（Sonnet/GLM 自动切换）
model_router = None
usage_tracker = None
try:
    usage_tracker = UsageTracker(DATABASE_URL)
    model_router = ModelRouter(usage_tracker)
    print("✓ 模型路由器初始化成功")
except Exception as e:
    print(f"⚠ 模型路由器初始化失败: {e}")

# LLM 服务（评估引擎用）- 使用 GLM 模型
llm_service = None
llm_service_pk = None
if model_router and model_router.glm_client:
    try:
        llm_service = LLMService(
            client=model_router.glm_client,
            model=model_router.glm_model
        )
        print(f"✓ LLM服务初始化成功（GLM 模型: {model_router.glm_model}）")
        llm_service_pk = llm_service
    except Exception as e:
        print(f"✗ GLM LLM服务初始化失败: {e}")

if llm_service is None:
    print("✗ LLM服务初始化失败: GLM 未配置，请设置 GLM_API_KEY 环境变量")
    exit(1)

# 简历优化 Agent
chat_agent = None
try:
    chat_agent = ChatAgent(
        client=llm_service.client,
        model=llm_service.model,
        llm_service=llm_service,
        convergence_engine=None,  # 在 convergence_engine 初始化后注入
        model_router=model_router,
    )
    print("✓ 简历优化Agent初始化成功")
except Exception as e:
    print(f"✗ 简历优化Agent初始化失败: {e}")

# 增量收敛引擎
convergence_engine = None
try:
    convergence_engine = IncrementalConvergence(
        validation_rules=validation_rules,
        llm_service=llm_service_pk
    )
    print("✓ 增量收敛引擎初始化成功")
    # 注入到 chat_agent（延迟注入，因为 chat_agent 先于 convergence_engine 初始化）
    if chat_agent:
        chat_agent.convergence_engine = convergence_engine
        print("  → 已注入 chat_agent（Function Call 工具调用已启用）")
except Exception as e:
    print(f"✗ 增量收敛引擎初始化失败: {e}")

# 薪酬计算器
salary_calculator = None
try:
    salary_calculator = SalaryCalculator()
    print("✓ 薪酬计算器初始化成功")
except Exception as e:
    print(f"✗ 薪酬计算器初始化失败: {e}")

print("=" * 60 + "\n")


# ===========================================
# 后端支持的职能类型（与前端选项一致）
# ===========================================

VALID_FUNCTIONS = {
    "算法", "软件开发", "产品管理", "数据分析与商业智能",
    "硬件开发", "信息安全", "投融资管理", "战略管理",
    "法务", "人力资源", "资产管理", "市场营销",
    "销售", "硬件测试", "税务", "内审",
    "软件测试", "产品运营", "公共关系", "游戏设计",
    "项目管理", "电商运营", "风险管理", "财务管理",
    "会计", "网络教育", "供应链管理", "广告",
    "采购", "客户服务", "物流", "行政管理",
    "IT服务", "销售运营", "媒体推广运营", "通用职能"
}

# 默认职能（当用户选择无效值时使用）
DEFAULT_FUNCTION = "行政管理"


# ===========================================
# API 路由
# ===========================================

@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查"""
    return jsonify({
        'status': 'ok',
        'service': 'student-value-backend',
        'llm': llm_service is not None,
        'convergence': convergence_engine is not None,
        'salary': salary_calculator is not None,
        'chat_agent': chat_agent is not None
    }), 200


@app.route('/chat')
def chat_page():
    """简历优化助手聊天页面"""
    return send_from_directory('static', 'chat.html')


ADMIN_LOG_KEY = os.getenv('ADMIN_LOG_KEY', '')

@app.route('/api/mini/logs', methods=['GET'])
def get_logs():
    """返回评估日志（JSON），支持 ?after=N 增量拉取"""
    key = request.args.get('key', '')
    if not ADMIN_LOG_KEY or key != ADMIN_LOG_KEY:
        return jsonify({'error': 'unauthorized'}), 401

    after = int(request.args.get('after', 0))
    new_logs = _query_logs(after)
    return jsonify({'boot_id': 'db', 'logs': new_logs}), 200


@app.route('/api/mini/verify-invite', methods=['POST'])
def verify_invite():
    """
    验证邀请码（一次性使用）

    请求体: { "inviteCode": "STAR2026" }
    成功: { "success": true }
    失败: { "success": false, "error": "..." }
    """
    data = request.get_json() or {}
    code = (data.get('inviteCode') or '').strip().upper()

    if not code:
        return jsonify({'success': False, 'error': '请输入邀请码'}), 400

    # 永久邀请码：直接通过，不消耗
    if code in PERMANENT_INVITE_CODES:
        print(f"[邀请码] 永久码 {code} 验证通过")
        return jsonify({'success': True}), 200

    # 一次性邀请码
    if code not in INVITE_CODES_ALL:
        return jsonify({'success': False, 'error': '邀请码无效'}), 400

    with _codes_lock:
        # 已消耗（评估完成）的码彻底失效
        if code in CONSUMED_INVITE_CODES:
            return jsonify({'success': False, 'error': '该邀请码已被使用'}), 400

        if code in USED_INVITE_CODES:
            return jsonify({'success': False, 'error': '该邀请码已被使用'}), 400

        # 标记为已使用
        USED_INVITE_CODES.add(code)
        _save_used_codes(USED_INVITE_CODES)

    print(f"[邀请码] 一次性码 {code} 验证通过并已消耗，剩余: {INVITE_CODES_ALL - USED_INVITE_CODES}")
    return jsonify({'success': True}), 200


@app.route('/api/mini/assess', methods=['POST'])
def assess():
    """
    小程序评估接口（仅支持 CV 模式）

    请求体:
    {
        "assessmentType": "CV",
        "city": "上海",
        "industry": "互联网/科技",
        "jobTitle": "产品经理",
        "jobFunction": "产品管理",
        "resumeText": "简历内容..."
    }

    响应:
    {
        "success": true,
        "data": {
            "salaryRange": "25万-35万",
            "level": 14,
            "levelTag": "业务脊梁",
            "levelDesc": "你是团队的中流砥柱...",
            "abilities": {...},        // 8能力详情
            "radarData": {...},        // 雷达图数据
            "abilitySummary": "..."    // 能力总结
        }
    }
    """
    import time
    global assessment_counter
    start_time = time.time()

    try:
        data = request.get_json()

        # 城市 → 城市等级映射
        CITY_TIER_MAP = {
            "北京": "一线城市", "上海": "一线城市", "深圳": "一线城市", "广州": "一线城市",
            "杭州": "二线城市", "南京": "二线城市", "成都": "二线城市", "武汉": "二线城市",
            "苏州": "二线城市", "西安": "二线城市",
            "其他": "三线城市",
        }

        # 参数提取
        assessment_type = 'CV'  # 仅支持 CV 模式
        raw_city = data.get('city', '上海')
        city = CITY_TIER_MAP.get(raw_city, '二线城市')
        industry = data.get('industry', '互联网')
        job_title = data.get('jobTitle', '')
        job_function_raw = data.get('jobFunction', '其他')
        resume_text_raw = data.get('resumeText', '')

        # 学生版新增字段
        school_name = (data.get('schoolName') or '').strip()
        education_level = (data.get('educationLevel') or '本科').strip()
        major = (data.get('major') or '').strip()
        company_type = (data.get('companyType') or '').strip()
        target_company = (data.get('targetCompany') or '').strip()
        school_tier = identify_school_tier(school_name)
        print(f"[学生版] 学校={school_name}, 专业={major}, 学历={education_level}, 层级={school_tier}")

        # 处理简历内容（支持 PDF 文件解析）
        resume_text = extract_resume_text(resume_text_raw)

        # 职能验证（前后端选项已一致，直接使用）
        job_function = job_function_raw if job_function_raw in VALID_FUNCTIONS else DEFAULT_FUNCTION

        # 参数校验
        if not job_title:
            return jsonify({'success': False, 'error': '请填写职位名称'}), 400

        if not resume_text:
            return jsonify({'success': False, 'error': '请提供简历内容'}), 400

        # 读取 retryCount（前端传来，用于信息不足二次容错）
        retry_count = int(data.get('retryCount', 0))

        # 读取前端传来的页面停留时长（秒）
        welcome_s = data.get('welcomeS')
        form_s = data.get('formS')

        # 评估文本
        eval_text = resume_text

        print(f"\n[小程序评估] 类型={assessment_type}, 职位={job_title}, 职能={job_function}, retryCount={retry_count}")
        print(f"[调试] 原始resume_text长度={len(resume_text_raw)}字符")
        print(f"[调试] 解析后resume_text长度={len(resume_text)}字符")
        print(f"[调试] eval_text前500字:\n{eval_text[:500] if eval_text else '(空)'}\n")

        # ===========================================
        # 核心评估流程
        # ===========================================

        # 1. 增量收敛引擎分析
        print("[步骤1] 增量收敛分析...")
        convergence_result = convergence_engine.find_optimal_solution(
            eval_text=eval_text,
            title=job_title,
            function=job_function,
            revenue_contribution={'type': 'not_quantifiable'},
            assessment_type=assessment_type
        )

        # 信息不足拦截：必须在 best_solution 检查之前，因为信息不足可能导致收敛失败
        if convergence_result and convergence_result.get('insufficient_input') and retry_count == 0:
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            log_id = _insert_log({
                'time': now, 'type': 'insufficient',
                'city': raw_city, 'industry': industry, 'function': job_function, 'title': job_title,
                'resume_len': len(resume_text),
                'invite_code': None,
                'welcome_s': welcome_s,
                'form_s': form_s,
                'school_name': school_name or None,
                'education_level': education_level or None,
                'school_tier': school_tier or None,
                'major': major or None,
                'company_type': company_type or None,
                'target_company': target_company or None,
            })
            assessment_counter = log_id
            print(f"\n⚠️ 信息不足拦截 #{log_id} | {now}")
            print(f"输入: {raw_city} | {industry} | {job_function} | {job_title}")
            print(f"简历: {len(resume_text)}字 | 首次提交，返回422")
            return jsonify({'success': False, 'error': 'insufficient_input'}), 422

        if not convergence_result or not convergence_result.get('best_solution'):
            # 如果收敛失败，使用默认值
            print("[警告] 收敛失败，使用默认评估")
            return _fallback_assessment(job_title, city, industry, job_function, school_name, education_level, school_tier)

        best_solution = convergence_result['best_solution']

        # 2. 提取 HAY 8因素
        factors = {
            'practical_knowledge': best_solution.get('practical_knowledge', 'D'),
            'managerial_knowledge': best_solution.get('managerial_knowledge', 'I'),
            'communication': best_solution.get('communication', '2'),
            'thinking_environment': best_solution.get('thinking_environment', 'D'),
            'thinking_challenge': best_solution.get('thinking_challenge', '3'),
            'freedom_to_act': best_solution.get('freedom_to_act', 'C'),
            'magnitude': best_solution.get('magnitude', 'N'),
            'nature_of_impact': best_solution.get('nature_of_impact', 'III')
        }

        # 3. 计算 HAY 评分和职级
        print("[步骤2] 计算HAY评分...")
        hay_result = calculate_hay_evaluation(factors)
        job_grade = hay_result['summary'].get('job_grade', 14)
        total_score = hay_result['summary'].get('total_score', 0)

        # 学生版：职级下限兜底到 9
        if job_grade < 9:
            job_grade = 9

        # 4. 查询薪酬
        print("[步骤3] 查询薪酬...")
        salary_result = None
        if salary_calculator:
            try:
                salary_result = salary_calculator.get_salary_range(
                    job_grade=job_grade,
                    function=job_function,
                    industry=industry,
                    city=city
                )
            except Exception as e:
                print(f"[警告] 薪酬查询失败: {e}")

        # 格式化薪酬区间（学生版：应用学校+学历系数，输出 k 格式）
        if salary_result:
            base_low = salary_result['P50_low']
            base_high = salary_result['P50_high']
        else:
            # 根据职级估算
            base_low = 8 + (job_grade - 10) * 3
            base_high = base_low + 10

        adj_low, adj_high = apply_student_coefficients(base_low, base_high, school_tier, education_level)
        salary_range = format_salary_k(adj_low, adj_high)
        print(f"[学生版] 薪酬: 基础{base_low:.0f}-{base_high:.0f}万/年 × {school_tier}/{education_level} → 月{salary_range}")

        # 5. 8因素 → 8能力维度映射
        print("[步骤4] 能力映射...")
        abilities = map_factors_to_dimensions(factors)
        radar_data = get_dimension_radar_data(abilities)
        ability_summary = get_dimension_summary(abilities)

        # 6. 生成趣味标签（学生版：传入 total_score 用于子档判定）
        print("[步骤5] 生成趣味标签...")
        level_tag, level_desc = get_level_tag_and_desc(job_grade, factors, abilities, total_score=total_score)

        # 7. AI 深度评估已移除（功能由聊天 Agent 承接）
        is_insufficient = bool(convergence_result and convergence_result.get('insufficient_input'))

        # 计算薪酬竞争力百分位
        salary_competitiveness = calculate_salary_competitiveness(job_function, job_grade)
        print(f"[步骤7] 薪酬竞争力: 超过{salary_competitiveness}%的同职能从业者")

        # ===========================================
        # 详细日志输出（验证映射逻辑）
        # ===========================================
        print("\n" + "=" * 60)
        print("【评估结果详情】")
        print("=" * 60)
        print(f"岗位职级: {job_grade}")
        print(f"薪酬范围: {salary_range}")
        print(f"趣味标签: {level_tag}")
        print("-" * 60)
        print("【HAY 8因素档位】")
        print(f"  Know-How (知识技能):")
        print(f"    - PK  专业知识:     {factors['practical_knowledge']}")
        print(f"    - MK  管理知识:     {factors['managerial_knowledge']}")
        print(f"    - Comm 沟通技巧:    {factors['communication']}")
        print(f"  Problem Solving (解决问题):")
        print(f"    - TE  思维环境:     {factors['thinking_environment']}")
        print(f"    - TC  思维挑战:     {factors['thinking_challenge']}")
        print(f"  Accountability (责任):")
        print(f"    - FTA 行动自由:     {factors['freedom_to_act']}")
        print(f"    - M   影响范围:     {factors['magnitude']}")
        print(f"    - NI  影响性质:     {factors['nature_of_impact']}")
        print("-" * 60)
        print("【8能力维度得分】")
        for name, info in abilities.items():
            print(f"  {name}: {info['score']}分 ({info['level']}) [{info.get('grade', '')}]")
        print("=" * 60 + "\n")

        # ===========================================
        # 返回结果
        # ===========================================

        # 结构化使用记录摘要
        elapsed_time = time.time() - start_time
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        log_id = _insert_log({
            'time': now, 'type': 'success',
            'city': raw_city, 'industry': industry, 'function': job_function, 'title': job_title,
            'resume_len': len(resume_text), 'retry': retry_count, 'insufficient': is_insufficient,
            'resume_text': resume_text,
            'grade': job_grade, 'tag': level_tag, 'salary': salary_range,
            'factors_detail': {
                'PK(专业知识)': factors['practical_knowledge'],
                'MK(管理知识)': factors['managerial_knowledge'],
                'Comm(沟通技巧)': factors['communication'],
                'TE(思维环境)': factors['thinking_environment'],
                'TC(思维挑战)': factors['thinking_challenge'],
                'FTA(行动自由)': factors['freedom_to_act'],
                'M(影响范围)': factors['magnitude'],
                'NI(影响性质)': factors['nature_of_impact'],
            },
            'abilities': " ".join(f"{name}{info['score']}" for name, info in abilities.items()),
            'deep_insight': None,  # 已移除，深度洞察功能由聊天 Agent 承接
            'elapsed': f"{elapsed_time:.2f}",
            'invite_code': None,
            'welcome_s': welcome_s,
            'form_s': form_s,
            'school_name': school_name or None,
            'education_level': education_level or None,
            'school_tier': school_tier or None,
        })
        assessment_counter = log_id
        print(f"\n══════════════════════════════════════════")
        print(f"📋 评估记录 #{log_id} | {now}")
        print(f"──────────────────────────────────────────")
        print(f"输入: {raw_city} | {industry} | {job_function} | {job_title}")
        print(f"学校: {school_name} ({school_tier}) | 学历: {education_level}")
        print(f"简历: {len(resume_text)}字 | 重试: {retry_count} | 信息不足: {'是' if is_insufficient else '否'}")
        print(f"──────────────────────────────────────────")
        print(f"结果: 职级{job_grade} | {level_tag} | {salary_range}")
        print(f"因素: PK={factors['practical_knowledge']} MK={factors['managerial_knowledge']} Comm={factors['communication']} TE={factors['thinking_environment']} TC={factors['thinking_challenge']}")
        abilities_str = ' '.join(f"{n}{d['score']}" for n, d in abilities.items())
        print(f"能力: {abilities_str}")
        print(f"耗时: {elapsed_time:.2f}秒")
        print(f"══════════════════════════════════════════\n")

        response_data = {
            # === 能力评估（纯简历驱动） ===
            'level': job_grade,
            'levelTag': level_tag,
            'levelDesc': level_desc,
            'abilities': abilities,
            'radarData': radar_data,
            'abilitySummary': ability_summary,
            'abilityCompetitiveness': salary_competitiveness,  # 能力百分位 0-100（基于职级在同届中的排名）
            'resumeHealthScore': _calc_resume_health(abilities),

            # === 市场薪酬参考（城市/行业/职能驱动） ===
            'salaryRange': salary_range,  # 保留以兼容旧前端
            'marketSalary': {
                'range': salary_range,
                'note': '月度基本工资（不含年终奖金及其他福利）',
                'city': raw_city,
                'industry': industry,
                'function': job_function,
            },
            'salaryCompetitiveness': salary_competitiveness,  # 向后兼容

            # 学生版附加信息
            'schoolTier': school_tier,

            # 解析后的简历文本（供聊天 Agent 使用）
            'resumeText': resume_text,

            # 调试信息（可选，生产环境可移除）
            'factors': factors,
            'logId': log_id,
        }

        # 如果前端传了 userId，同步存入 Supabase assessments 表
        user_id = data.get('userId')
        if user_id and supabase_client:
            try:
                supabase_client.table('assessments').insert({
                    'user_id': user_id,
                    'resume_text': resume_text,
                    'form_data': {
                        'city': raw_city, 'industry': industry,
                        'jobTitle': job_title, 'jobFunction': job_function,
                        'educationLevel': education_level, 'major': major,
                        'companyType': company_type, 'targetCompany': target_company,
                    },
                    'result': response_data,
                }).execute()
                print(f"[Supabase] 评估记录已存入 assessments 表 (user={user_id[:8]}...)")
            except Exception as sb_err:
                print(f"[Supabase] 评估存储失败: {sb_err}")

        return jsonify({'success': True, 'data': response_data}), 200

    except Exception as e:
        elapsed_time = time.time() - start_time
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        # 安全提取输入信息（data 可能解析失败）
        _city = locals().get('raw_city', '?')
        _industry = locals().get('industry', '?')
        _func = locals().get('job_function', locals().get('job_function_raw', '?'))
        _title = locals().get('job_title', '?')
        log_id = _insert_log({
            'time': now, 'type': 'error',
            'city': _city, 'industry': _industry, 'function': _func, 'title': _title,
            'error': f"{type(e).__name__}: {e}", 'elapsed': f"{elapsed_time:.2f}",
            'invite_code': None,
            'welcome_s': locals().get('welcome_s'),
            'form_s': locals().get('form_s'),
        })
        assessment_counter = log_id
        print(f"\n❌ 评估失败 #{log_id} | {now}")
        print(f"输入: {_city} | {_industry} | {_func} | {_title}")
        print(f"错误: {type(e).__name__}: {e}")
        print(f"耗时: {elapsed_time:.2f}秒")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def _fallback_assessment(job_title, city, industry, job_function, school_name='', education_level='本科', school_tier='普通本科'):
    """
    降级评估（学生版，当主流程失败时使用）
    """
    # 学生版默认职级（下移）
    default_grades = {
        "算法": 12, "软件开发": 11, "产品管理": 11, "数据分析与商业智能": 11,
        "硬件开发": 11, "信息安全": 10, "投融资管理": 11, "战略管理": 11,
        "法务": 10, "人力资源": 10, "资产管理": 10, "市场营销": 10,
        "销售": 10, "硬件测试": 10, "税务": 10, "内审": 10,
        "软件测试": 10, "产品运营": 10, "公共关系": 9, "游戏设计": 9,
        "项目管理": 10, "电商运营": 9, "风险管理": 9, "财务管理": 10,
        "会计": 9, "网络教育": 9, "供应链管理": 9, "广告": 9,
        "采购": 9, "客户服务": 9, "物流": 9, "行政管理": 9,
        "IT服务": 9, "销售运营": 9, "媒体推广运营": 9, "通用职能": 10
    }
    job_grade = default_grades.get(job_function, 10)

    # 城市系数
    city_factor = 1.0
    if city in ["上海", "北京", "深圳"]:
        city_factor = 1.2
    elif city in ["杭州", "广州"]:
        city_factor = 1.1

    base_low = (8 + (job_grade - 10) * 3) * city_factor
    base_high = base_low + 10
    adj_low, adj_high = apply_student_coefficients(base_low, base_high, school_tier, education_level)
    salary_range = format_salary_k(adj_low, adj_high)

    # 默认能力（5维度）
    default_abilities = {
        "专业力": {"score": 55, "level": "medium", "explanation": "具备扎实的专业基础，能够独立完成常规专业工作"},
        "管理力": {"score": 45, "level": "medium", "explanation": "能够管理自己的工作任务，配合团队完成目标"},
        "合作力": {"score": 50, "level": "medium", "explanation": "能够在团队内部有效沟通，配合完成协作任务"},
        "思辨力": {"score": 50, "level": "medium", "explanation": "能够按照既定框架分析和解决问题"},
        "创新力": {"score": 45, "level": "medium", "explanation": "能够在现有框架下完成工作，学习新事物"},
    }

    return jsonify({
        'success': True,
        'data': {
            'level': job_grade,
            'levelTag': "萌新探路者",
            'levelDesc': "刚踏出校门第一步，世界很大，你的好奇心更大。起点不决定终点，你的故事才刚刚开始写呢。",
            'abilities': default_abilities,
            'radarData': {name: info["score"] for name, info in default_abilities.items()},
            'abilitySummary': "建议持续提升专业能力，拓展校招竞争力。",
            'abilityCompetitiveness': 30,
            'resumeHealthScore': 40,
            'salaryRange': salary_range,
            'marketSalary': {
                'range': salary_range,
                'note': '月度基本工资（不含年终奖金及其他福利）',
                'city': city,
                'industry': industry,
                'function': job_function,
            },
            'salaryCompetitiveness': 30,
            'schoolTier': school_tier,
            'factors': None,
        }
    }), 200


VALID_DURATION_COLS = {'result_s', 'mine_s'}

@app.route('/api/mini/update-duration', methods=['POST'])
def update_duration():
    """回写评估后页面的停留时长到 assessment_logs"""
    data = request.get_json() or {}
    log_id = data.get('logId')
    col = data.get('column')  # result_s / mine_s
    duration_s = data.get('durationS')

    if not log_id or not isinstance(log_id, int):
        return jsonify({'success': False, 'error': 'invalid logId'}), 400
    if col not in VALID_DURATION_COLS:
        return jsonify({'success': False, 'error': 'invalid column'}), 400
    if not isinstance(duration_s, (int, float)) or duration_s <= 0:
        return jsonify({'success': False, 'error': 'invalid durationS'}), 400

    if not DATABASE_URL:
        return jsonify({'success': True}), 200
    try:
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor()
        cur.execute(f'UPDATE assessment_logs SET {col} = %s WHERE id = %s', (int(duration_s), log_id))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"[页面停留] ❌ 更新失败: {e}")
    return jsonify({'success': True}), 200


# ===========================================
# 简历优化 Agent API
# ===========================================


def _parse_edit_block(edit_block: str, resume_sections: list) -> dict:
    """
    解析 <<<EDIT...EDIT>>> 块，匹配 section

    Returns:
        {"sectionId": "section-0", "original": "...", "suggested": "...", "rationale": "..."}
        或 None（解析失败）
    """
    try:
        # 提取各字段
        section_match = re.search(r'SECTION:\s*(.+?)(?:\n|$)', edit_block)
        original_match = re.search(r'ORIGINAL:\s*(.+?)(?=\nSUGGESTED:)', edit_block, re.DOTALL)
        suggested_match = re.search(r'SUGGESTED:\s*(.+?)(?=\nRATIONALE:)', edit_block, re.DOTALL)
        rationale_match = re.search(r'RATIONALE:\s*(.+?)(?=\nEDIT>>>)', edit_block, re.DOTALL)

        if not all([section_match, original_match, suggested_match, rationale_match]):
            print(f"[Agent API] 编辑块解析失败，字段不完整")
            return None

        section_title = section_match.group(1).strip()
        original = original_match.group(1).strip()
        suggested = suggested_match.group(1).strip()
        rationale = rationale_match.group(1).strip()

        # 按 title 模糊匹配 section
        matched_id = None
        for i, sec in enumerate(resume_sections):
            sec_title = sec.get('title', '')
            if section_title in sec_title or sec_title in section_title:
                matched_id = f'section-{i}'
                break

        # 如果没精确匹配，尝试部分匹配
        if not matched_id:
            for i, sec in enumerate(resume_sections):
                sec_title = sec.get('title', '')
                # 取标题中的关键词匹配
                if any(kw in sec_title for kw in section_title.split('-') if len(kw) >= 2):
                    matched_id = f'section-{i}'
                    break

        if not matched_id and resume_sections:
            # 最后兜底：匹配内容包含 original 的 section
            for i, sec in enumerate(resume_sections):
                if original[:20] in sec.get('content', ''):
                    matched_id = f'section-{i}'
                    break

        if not matched_id:
            matched_id = 'section-0'  # 兜底

        return {
            'sectionId': matched_id,
            'original': original,
            'suggested': suggested,
            'rationale': rationale,
        }
    except Exception as e:
        print(f"[Agent API] 编辑块解析异常: {e}")
        return None


@app.route('/api/chat/start', methods=['POST'])
def chat_start():
    """
    开启简历优化对话

    请求体:
    {
        "assessmentContext": {
            "factors": {...},         // HAY 8因素
            "abilities": {...},       // 8维能力
            "grade": 12,
            "salaryRange": "8.5k~12k",
            "jobTitle": "产品经理",
            "jobFunction": "产品管理",
            "deepInsight": "..."
        },
        "resumeText": "简历原文..."
    }

    响应:
    {
        "success": true,
        "data": {
            "sessionId": "uuid",
            "greeting": "你好！我看了你的评测结果..."
        }
    }
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    try:
        data = request.get_json()
        assessment_context = data.get('assessmentContext', {})
        resume_text_raw = data.get('resumeText', '')

        if not assessment_context:
            return jsonify({'success': False, 'error': '缺少评测上下文'}), 400

        if not resume_text_raw:
            return jsonify({'success': False, 'error': '缺少简历内容'}), 400

        # 解析简历（复用现有的文件解析逻辑）
        resume_text = extract_resume_text(resume_text_raw)
        if not resume_text:
            return jsonify({'success': False, 'error': '简历内容解析失败'}), 400

        # 创建会话并生成开场白
        user_id = data.get('userId')
        assessment_id = data.get('assessmentId')
        result = chat_agent.start_session(
            assessment_context=assessment_context,
            resume_text=resume_text,
            user_id=user_id,
            assessment_id=assessment_id,
        )

        print(f"[Agent API] 会话已创建: {result['session_id'][:8]}...")

        return jsonify({
            'success': True,
            'data': {
                'sessionId': result['session_id'],
                'greeting': result['greeting']
            }
        }), 200

    except Exception as e:
        print(f"[Agent API] /chat/start 失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/chat/message', methods=['POST'])
def chat_message():
    """
    发送消息（支持普通模式和流式模式）

    请求体:
    {
        "sessionId": "uuid",
        "message": "帮我改一下实习经历那段",
        "stream": true   // 可选，默认 false
    }

    普通模式响应:
    {
        "success": true,
        "data": {
            "reply": "好的，我来看看你的实习经历..."
        }
    }

    流式模式响应: SSE (text/event-stream)
        data: {"type": "text", "content": "好的"}
        data: {"type": "text", "content": "，我来"}
        ...
        data: {"type": "done"}
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    try:
        data = request.get_json()
        session_id = data.get('sessionId', '')
        message = data.get('message', '').strip()
        use_stream = data.get('stream', False)

        if not session_id:
            return jsonify({'success': False, 'error': '缺少 sessionId'}), 400
        if not message:
            return jsonify({'success': False, 'error': '消息不能为空'}), 400

        # 消息长度限制（防止 token 超限）
        MAX_MESSAGE_LEN = 2000
        if len(message) > MAX_MESSAGE_LEN:
            message = message[:MAX_MESSAGE_LEN]
            print(f"[Agent API] 用户消息过长，已截断至 {MAX_MESSAGE_LEN} 字符")

        # 验证会话存在
        session = chat_agent.session_manager.get_session(session_id)
        if not session:
            return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

        canvas_mode = data.get('canvasMode', False)

        if use_stream:
            # 流式模式：SSE
            def generate():
                try:
                    if canvas_mode:
                        # 画布模式：缓冲检测 <<<EDIT...EDIT>>> 块
                        buffer = ''
                        session = chat_agent.session_manager.get_session(session_id)
                        resume_sections = session.get('resume_sections', []) if session else []

                        for chunk in chat_agent.chat_stream(session_id, message, canvas_mode=True):
                            buffer += chunk

                            # 持续处理缓冲区中的完整编辑块
                            while '<<<EDIT' in buffer and 'EDIT>>>' in buffer:
                                edit_start = buffer.index('<<<EDIT')
                                edit_end = buffer.index('EDIT>>>') + len('EDIT>>>')

                                # 发送编辑块之前的普通文本
                                before_text = buffer[:edit_start].strip()
                                if before_text:
                                    yield f"data: {json.dumps({'type': 'text', 'content': before_text}, ensure_ascii=False)}\n\n"

                                # 解析编辑块
                                edit_block = buffer[edit_start:edit_end]
                                edit_data = _parse_edit_block(edit_block, resume_sections)
                                if edit_data:
                                    yield f"data: {json.dumps({'type': 'edit', **edit_data}, ensure_ascii=False)}\n\n"

                                buffer = buffer[edit_end:]

                        # 发送剩余的普通文本
                        remaining = buffer.strip()
                        if remaining:
                            yield f"data: {json.dumps({'type': 'text', 'content': remaining}, ensure_ascii=False)}\n\n"
                    else:
                        # 普通模式
                        for chunk in chat_agent.chat_stream(session_id, message):
                            yield f"data: {json.dumps({'type': 'text', 'content': chunk}, ensure_ascii=False)}\n\n"

                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                except Exception as e:
                    print(f"[Agent API] 流式输出错误: {e}")
                    yield f"data: {json.dumps({'type': 'error', 'content': str(e)}, ensure_ascii=False)}\n\n"

            return Response(
                generate(),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'X-Accel-Buffering': 'no',
                    'Connection': 'keep-alive',
                }
            )
        else:
            # 普通模式
            reply = chat_agent.chat(session_id, message)
            if reply is None:
                return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

            return jsonify({
                'success': True,
                'data': {'reply': reply}
            }), 200

    except Exception as e:
        print(f"[Agent API] /chat/message 失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/chat/sections', methods=['GET'])
def chat_sections():
    """
    获取简历结构化段落

    参数: ?sessionId=uuid

    响应:
    - 解析中: {"success": true, "data": {"status": "parsing"}}
    - 完成:   {"success": true, "data": {"status": "ready", "sections": [...]}}
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    session_id = request.args.get('sessionId', '')
    if not session_id:
        return jsonify({'success': False, 'error': '缺少 sessionId'}), 400

    session = chat_agent.session_manager.get_session(session_id)
    if not session:
        return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

    resume_sections = session.get('resume_sections')
    if resume_sections is None:
        return jsonify({'success': True, 'data': {'status': 'parsing'}}), 200

    # 为每个 section 补充 id 字段
    sections_with_id = []
    for i, sec in enumerate(resume_sections):
        sections_with_id.append({
            'id': f'section-{i}',
            'type': sec.get('type', 'other'),
            'title': sec.get('title', ''),
            'content': sec.get('content', ''),
        })

    return jsonify({
        'success': True,
        'data': {'status': 'ready', 'sections': sections_with_id}
    }), 200


@app.route('/api/chat/edit-action', methods=['POST'])
def chat_edit_action():
    """
    处理用户对编辑建议的采纳/忽略操作

    请求体:
    {
        "sessionId": "uuid",
        "sectionId": "section-0",
        "action": "accept" | "reject",
        "suggestedText": "改写后的文本"
    }
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    try:
        data = request.get_json()
        session_id = data.get('sessionId', '')
        section_id = data.get('sectionId', '')
        action = data.get('action', '')
        suggested_text = data.get('suggestedText', '')

        if not session_id or not section_id or action not in ('accept', 'reject'):
            return jsonify({'success': False, 'error': '参数不完整'}), 400

        session = chat_agent.session_manager.get_session(session_id)
        if not session:
            return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

        resume_sections = session.get('resume_sections')
        if not resume_sections:
            return jsonify({'success': False, 'error': '简历段落数据不存在'}), 400

        # 从 sectionId 提取索引
        try:
            idx = int(section_id.replace('section-', ''))
        except ValueError:
            return jsonify({'success': False, 'error': '无效的 sectionId'}), 400

        if idx < 0 or idx >= len(resume_sections):
            return jsonify({'success': False, 'error': 'sectionId 超出范围'}), 400

        memory = session.get('memory')

        if action == 'accept':
            # 更新 section content
            resume_sections[idx]['content'] = suggested_text
            chat_agent.session_manager.update_session(session_id, {
                'resume_sections': resume_sections
            })
            # 记入结构化记忆
            if memory:
                title = resume_sections[idx].get('title', '')
                memory.add_agreed_modification(f"采纳了「{title}」段落的修改建议")
        else:
            # reject：记入 memory 避免重复建议
            if memory:
                title = resume_sections[idx].get('title', '')
                memory.add_agreed_modification(f"用户拒绝了「{title}」段落的修改建议，不要重复建议")

        return jsonify({'success': True, 'data': {'action': action, 'sectionId': section_id}}), 200

    except Exception as e:
        print(f"[Agent API] /chat/edit-action 失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/chat/history', methods=['GET'])
def chat_history():
    """
    获取对话历史

    参数: ?sessionId=uuid

    响应:
    {
        "success": true,
        "data": {
            "messages": [
                {"role": "user", "content": "..."},
                {"role": "assistant", "content": "..."},
                ...
            ]
        }
    }
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    session_id = request.args.get('sessionId', '')
    if not session_id:
        return jsonify({'success': False, 'error': '缺少 sessionId'}), 400

    history = chat_agent.get_history(session_id)
    if history is None:
        return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

    return jsonify({
        'success': True,
        'data': {'messages': history}
    }), 200


# ===========================================
# 用户评估历史 API（Supabase）
# ===========================================

@app.route('/api/user/assessments', methods=['GET'])
def user_assessments():
    """
    获取用户的评估历史列表

    参数: ?userId=uuid
    """
    user_id = request.args.get('userId', '')
    if not user_id:
        return jsonify({'success': False, 'error': '缺少 userId'}), 400

    if not supabase_client:
        return jsonify({'success': False, 'error': 'Supabase 未配置'}), 503

    try:
        resp = supabase_client.table('assessments') \
            .select('id, form_data, result, created_at') \
            .eq('user_id', user_id) \
            .order('created_at', desc=True) \
            .execute()
        return jsonify({'success': True, 'data': resp.data}), 200
    except Exception as e:
        print(f"[用户历史] 查询失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ===========================================
# 用户 LLM 用量查询 API
# ===========================================

@app.route('/api/user/usage', methods=['GET'])
def user_usage():
    """
    获取用户的 LLM 用量摘要

    参数: ?userId=uuid

    响应:
    {
        "success": true,
        "data": {
            "user_id": "...",
            "sonnet_cost_rmb": 3.5,
            "sonnet_budget_rmb": 15.0,
            "sonnet_remaining_rmb": 11.5,
            "budget_exceeded": false
        }
    }
    """
    user_id = request.args.get('userId', '')
    if not user_id:
        return jsonify({'success': False, 'error': '缺少 userId'}), 400

    if not usage_tracker:
        return jsonify({'success': False, 'error': '用量追踪未启用'}), 503

    try:
        summary = usage_tracker.get_user_usage_summary(user_id)
        return jsonify({'success': True, 'data': summary}), 200
    except Exception as e:
        print(f"[用量查询] 查询失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ===========================================
# 简历 Word 导出
# ===========================================

@app.route('/api/chat/resume/export', methods=['POST'])
def chat_resume_export():
    """
    将当前 session 的简历导出为 Word 文档

    请求体: { "sessionId": "uuid" }
    响应: .docx 文件流
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    try:
        data = request.get_json()
        session_id = data.get('sessionId', '')
        if not session_id:
            return jsonify({'success': False, 'error': '缺少 sessionId'}), 400

        session = chat_agent.session_manager.get_session(session_id)
        if not session:
            return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

        resume_sections = session.get('resume_sections')
        if not resume_sections:
            return jsonify({'success': False, 'error': '简历数据不存在，请先进入简历画布'}), 400

        from resume_export import generate_resume_docx
        buffer = generate_resume_docx(resume_sections)

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='我的简历.docx',
        )

    except Exception as e:
        print(f"[简历导出] 失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ===========================================
# 简历 PDF 导出
# ===========================================

@app.route('/api/chat/resume/export-pdf', methods=['POST'])
def chat_resume_export_pdf():
    """
    将当前 session 的简历导出为 PDF 文档

    请求体: { "sessionId": "uuid" }
    响应: .pdf 文件流
    """
    if not chat_agent:
        return jsonify({'success': False, 'error': 'Agent 服务未初始化'}), 503

    try:
        data = request.get_json()
        session_id = data.get('sessionId', '')
        if not session_id:
            return jsonify({'success': False, 'error': '缺少 sessionId'}), 400

        session = chat_agent.session_manager.get_session(session_id)
        if not session:
            return jsonify({'success': False, 'error': '会话不存在或已过期'}), 404

        resume_sections = session.get('resume_sections')
        if not resume_sections:
            return jsonify({'success': False, 'error': '简历数据不存在，请先进入简历画布'}), 400

        from resume_export import generate_resume_pdf
        buffer = generate_resume_pdf(resume_sections)

        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='我的简历.pdf',
        )

    except Exception as e:
        print(f"[简历PDF导出] 失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ===========================================
# 启动服务
# ===========================================

if __name__ == '__main__':
    port = 5001
    print(f"\n学生版校招身价评估后端服务启动: http://localhost:{port}")
    print("=" * 60 + "\n")
    app.run(debug=True, host='0.0.0.0', port=port)
